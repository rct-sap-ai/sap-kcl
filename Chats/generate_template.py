import os
import sys
import glob
import json
import re
import zipfile
from xml.etree import ElementTree as ET
from typing import Dict, Optional, Tuple


def find_template_path(root_dir: str) -> str:
    """Locate the DOCX template file.

    Tries common locations first, then falls back to a recursive search.
    Returns absolute path or raises FileNotFoundError.
    """
    candidate_names = [
        os.path.join(root_dir, "Templates", "DRAFT Q-162 (SAP Template) V2.0.docx"),
        os.path.join(root_dir, "Templates SAP", "DRAFT Q-162 (SAP Template) V2.0.docx"),
        os.path.join(root_dir, "Templates", "SAP", "DRAFT Q-162 (SAP Template) V2.0.docx"),
    ]
    for p in candidate_names:
        if os.path.isfile(p):
            return p

    # Fallback: recursive glob
    pattern = os.path.join(root_dir, "**", "DRAFT Q-162 (SAP Template) V2.0.docx")
    matches = glob.glob(pattern, recursive=True)
    if matches:
        return os.path.abspath(matches[0])

    raise FileNotFoundError(
        "Template file 'DRAFT Q-162 (SAP Template) V2.0.docx' not found under Templates."
    )


def find_sap_json(root_dir: str, override_path: Optional[str] = None) -> str:
    """Find the SAP JSON file to use for replacements.

    Priority order:
    1) override_path if provided and exists
    2) Environment variable SAP_JSON pointing to a file
    3) Newest *.json under root/SAP
    """
    if override_path and os.path.isfile(override_path):
        return os.path.abspath(override_path)

    env_path = os.getenv("SAP_JSON", "").strip()
    if env_path and os.path.isfile(env_path):
        return os.path.abspath(env_path)

    sap_dir = os.path.join(root_dir, "SAP")
    json_files = sorted(
        glob.glob(os.path.join(sap_dir, "*.json")),
        key=lambda p: os.path.getmtime(p),
        reverse=True,
    )
    if not json_files:
        raise FileNotFoundError(f"No JSON files found in: {sap_dir}")
    return os.path.abspath(json_files[0])


def load_replacements(json_path: str) -> Dict[str, str]:
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    # Ensure values are strings
    return {str(k): ("" if v is None else str(v)) for k, v in data.items()}


def _prepare_tag_name_map(replacements: Dict[str, str]) -> Dict[str, str]:
    """Normalize JSON keys into tag names without braces/spaces, lowercased.

    Examples:
      "{{title}}"           -> name "title"
      "{{ title }}"        -> name "title"
      "{{Stratification}}" -> name "stratification"
      "title"              -> name "title" (fallback)
    """
    name_map: Dict[str, str] = {}
    for key, value in replacements.items():
        m = re.fullmatch(r"\{\{\s*([^}]+?)\s*\}\}", key)
        if m:
            name = m.group(1).strip().lower()
        else:
            # fallback: treat the whole key as the name
            name = key.strip().strip("{} ").lower()
        name_map[name] = value
    return name_map


def _build_regex_patterns(name_map: Dict[str, str]) -> Tuple[Tuple[re.Pattern, str], ...]:
    """Build case/whitespace-insensitive regex patterns for each tag name.

    Matches any of the following: {{name}}, {{ name }}, {{NAME}}, etc.
    """
    patterns = []
    for name, value in name_map.items():
        pat = re.compile(r"\{\{\s*" + re.escape(name) + r"\s*\}\}", re.IGNORECASE)
        patterns.append((pat, value))
    return tuple(patterns)


def _replace_text(text: str, replacements: Dict[str, str], name_map: Optional[Dict[str, str]] = None, patterns: Optional[Tuple[Tuple[re.Pattern, str], ...]] = None) -> str:
    # 1) Exact string replacements first (for perfectly matching keys)
    for key, value in replacements.items():
        if key in text:
            text = text.replace(key, value)
    # 2) Regex replacements to catch {{ name }} variants and case differences
    if name_map is not None and patterns is not None:
        for pat, value in patterns:
            text = pat.sub(value, text)
    return text


def replace_placeholders_in_doc(doc, replacements: Dict[str, str]) -> None:
    """Replace placeholders in paragraphs and tables. This will simplify
    inline formatting (placeholders spanning multiple runs will be merged)."""
    from docx.text.paragraph import Paragraph
    from docx.table import _Cell

    # Prepare flexible matching of tag names
    name_map = _prepare_tag_name_map(replacements)
    patterns = _build_regex_patterns(name_map)

    # Body paragraphs
    for para in list(doc.paragraphs):
        if not para.text:
            continue
        new_text = _replace_text(para.text, replacements, name_map, patterns)
        if new_text != para.text:
            # Reset paragraph text (may lose inline run formatting)
            para.text = new_text

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if not para.text:
                        continue
                    new_text = _replace_text(para.text, replacements, name_map, patterns)
                    if new_text != para.text:
                        para.text = new_text

    # Headers and Footers (paragraphs and tables)
    for section in doc.sections:
        # Header
        header = section.header
        for para in header.paragraphs:
            if para.text:
                new_text = _replace_text(para.text, replacements, name_map, patterns)
                if new_text != para.text:
                    para.text = new_text
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text:
                            new_text = _replace_text(para.text, replacements, name_map, patterns)
                            if new_text != para.text:
                                para.text = new_text
        # Footer
        footer = section.footer
        for para in footer.paragraphs:
            if para.text:
                new_text = _replace_text(para.text, replacements, name_map, patterns)
                if new_text != para.text:
                    para.text = new_text
        for table in footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text:
                            new_text = _replace_text(para.text, replacements, name_map, patterns)
                            if new_text != para.text:
                                para.text = new_text


def generate_sap_doc(template_path: str, replacements: Dict[str, str], output_path: str) -> None:
    """Generate SAP.docx by performing placeholder replacement at the XML level.

    This approach handles text inside text boxes/shapes and other containers that
    python-docx can't reach. It merges runs by rewriting each paragraph's text.
    """

    # Prepare replacement helpers
    name_map = _prepare_tag_name_map(replacements)
    patterns = _build_regex_patterns(name_map)

    # Namespaces
    NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    NS_XML = "http://www.w3.org/XML/1998/namespace"

    # Ensure namespace prefixes are preserved on write
    ET.register_namespace('w', NS_W)

    def process_word_xml(xml_bytes: bytes) -> bytes:
        try:
            root = ET.fromstring(xml_bytes)
        except ET.ParseError:
            # Not a typical WordprocessingML part, return as-is
            return xml_bytes

        # Find all paragraphs anywhere in this part
        for p in root.findall('.//{'+NS_W+'}p'):
            # Gather all text nodes in this paragraph
            ts = p.findall('.//{'+NS_W+'}t')
            if not ts:
                continue
            original = ''.join([t.text or '' for t in ts])
            if not original:
                continue

            new_text = _replace_text(original, replacements, name_map, patterns)
            if new_text == original:
                continue

            # Clear paragraph children and write a single run with new_text
            for child in list(p):
                p.remove(child)
            r = ET.SubElement(p, '{'+NS_W+'}r')
            t = ET.SubElement(r, '{'+NS_W+'}t')
            # Preserve spaces/newlines as-is
            t.set('{'+NS_XML+'}space', 'preserve')
            t.text = new_text

        return ET.tostring(root, encoding='utf-8', xml_declaration=True)

    # Copy template zip, processing Word XML parts
    with zipfile.ZipFile(template_path, 'r') as zin:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        with zipfile.ZipFile(output_path, 'w', compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename.startswith('word/') and item.filename.endswith('.xml'):
                    # Process all WordprocessingML parts (document, header, footer, etc.)
                    data = process_word_xml(data)
                zout.writestr(item, data)


def main(argv=None):
    argv = argv or sys.argv[1:]
    root_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))

    # Optional: allow passing a specific JSON path as first argument
    override_json = argv[0] if argv else None

    template_path = find_template_path(root_dir)
    json_path = find_sap_json(root_dir, override_path=override_json)
    replacements = load_replacements(json_path)

    # Save SAP.docx in the same folder as the template
    output_dir = os.path.dirname(template_path)
    output_path = os.path.join(output_dir, "SAP.docx")

    print(f"Using template: {template_path}")
    print(f"Using JSON:     {json_path}")
    print(f"Writing to:     {output_path}")

    generate_sap_doc(template_path, replacements, output_path)
    print("Done.")


if __name__ == "__main__":
    main()
